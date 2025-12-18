from docx import Document
from docx.shared import Inches,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

list1=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",1,2,3],[4,5,6,7,8,9,10],[11,12,13,14,15,16,17],[18,19,20,21,22,23,24],[25,26,27,28,29,30,31]]
list2=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[1,2,3,4,5,6,7],[8,9,10,11,12,13,14],[15,16,17,18,19,20,21],[22,23,24,25,26,27,28],[".",".",".",".",".",".",".",]]
list3=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[1,2,3,4,5,6,7],[8,9,10,11,12,13,14],[15,16,17,18,19,20,21],[22,23,24,25,26,27,28],[29,30,31,".",".",".","."]]
list4=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",1,2,3,4],[5,6,7,8,9,10,11],[12,13,14,15,16,17,18],[19,20,21,22,23,24,25],[26,27,28,29,30,".","."]]
list5=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",".",1,2],[3,4,5,6,7,8,9],[10,11,12,13,14,15,16],[17,18,19,20,21,22,23],[24,25,26,27,28,29,30],[31,".",".",".",".",".","."]]
list6=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",1,2,3,4,5,6],[7,8,9,10,11,12,13],[14,15,16,17,18,19,20],[21,22,23,24,25,26,27],[28,29,30,".",".",".","."]]
list7=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",1,2,3,4],[5,6,7,8,9,10,11],[12,13,14,15,16,17,18],[19,20,21,22,23,24,25],[26,27,28,29,30,31,"."]]
list8=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",".",".",1],[2,3,4,5,6,7,8],[9,10,11,12,13,14,15],[16,17,18,19,20,21,22],[23,24,25,26,27,28,29],[30,31,".",".",".",".",","]]
list9=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",1,2,3,4,5],[6,7,8,9,10,11,12],[13,14,15,16,17,18,19],[20,21,22,23,24,25,26],[27,28,29,30,31,".","."]]
list10=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",1,2,3],[4,5,6,7,8,9,10],[11,12,13,14,15,16,17],[18,19,20,21,22,23,24],[25,26,27,28,29,30,31]]
list11=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[1,2,3,4,5,6,7],[8,9,10,11,12,13,14],[15,16,17,18,19,20,21],[22,23,24,25,26,27,28],[29,30,".",".",".",".","."]]
list12=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",1,2,3,4,5],[6,7,8,9,10,11,12],[13,14,15,16,17,18,19],[20,21,22,23,24,25,26],[27,28,29,30,31,".","."]]
listfull=[list1,list2,list3,list4,list5,list6,list7,list8,list9,list10,list11,list12]
year="2026"
month=["January","February","March","April","May","June","July","August","September","October","November","December"]
holiday= [{"date": "15-01-2026","name": "Makara Sankranthi","Day": "Thursday"},
        {"date": "26-01-2026", "name": "Republic Day","Day": "Monday"},
        {"date": "19-03-2026", "name": "Chandramana Ugadi", "Day": "Thursday"},
        {"date": "27-03-2026", "name": "Shrirama Navami", "Day": "Friday"},
        {"date": "03-04-2026", "name": "Good Friday", "Day": "Friday"},
        {"date": "14-04-2026", "name": "Ambedkar Jayanthi", "Day": "Tuesday"},
        {"date": "01-05-2026", "name": "Labour day", "Day": "Friday"},
        {"date": "28-05-2026", "name": "Bakrid", "Day": "Thursday"},
        {"date": "26-06-2026", "name": "Moharam", "Day": "Friday"},
        {"date": "15-08-2026", "name": "Independence day", "Day": "Saturday"},
        {"date": "28-01-2026", "name": "Ed-Milad", "Day": "wednesday"},
        {"date": "05-09-2026", "name": "Vitla pindi", "Day": "Saturday"},
        {"date": "14-09-2026", "name": "Ganesha chathurthi", "Day": "Monday"},
        {"date": "02-10-2026", "name": "Gandi Jayanthi", "Day": "Friday"},
        {"date": "20-10-2026", "name": "Maha Navami / Ayudha pooja", "Day": "Tuesday"},
        {"date": "21-10-2026", "name": "Dassehra / vijayadashami", "Day": "wednesday"},
        {"date": "10-11-2026", "name": "Deepavali", "Day": "Tuesday"},
        {"date": "25-12-2026", "name": "Christmas", "Day": "Friday"}]
pic="tiger.jpg"
print("Done check")
count=0
doc=Document()
row=6
for listm in listfull:
    if(count<12):
        if(count==4 or count==7):
            row=7
        else:
            row=6
        doc.add_picture(pic)
        doc.add_paragraph("SMVITM")
        
        table1=doc.add_table(rows=1,cols=3)

        mon=table1.cell(0,0)
        run=mon.paragraphs[0].add_run(str(month[count]))
        run.font.color.rgb=RGBColor(0,0,255)

        year1=table1.cell(0,1)
        run = year1.paragraphs[0].add_run(str(year))
        run.font.color.rgb = RGBColor(128, 0, 0)

        table = doc.add_table(rows=row,cols=7)    
        table.style='Table Grid'

        table.columns[0].width, table.columns[1].width, table.columns[2].width, table.columns[3].width, table.columns[4].width, table.columns[5].width, table.columns[6].width =Inches(1.0), Inches(2.0), Inches(2.0), Inches(2.0), Inches(2.0), Inches(2.0), Inches(2.0), 
        table.rows[0].height, table.rows[1].height, table.rows[2].height, table.rows[3].height, table.rows[4].height, table.rows[5].height = Inches(0.5), Inches(1), Inches(1), Inches(1), Inches(1), Inches(1)

        for i in range(7):
            day0=table.cell(0, i)
            day0.text=str(listm[0][i])
            
        for i in range(7):
            day1=table.cell(1, i)
            day1.text=str(listm[1][i])
            
        for i in range(7):
            day2=table.cell(2, i)
            day2.text=str(listm[2][i])
            
        for i in range(7):
            day3=table.cell(3, i)
            day3.text=str(listm[3][i])

        for i in range(7):
            day4=table.cell(4, i)
            day4.text = str(listm[4][i])
            
        for i in range(7):
            day5=table.cell(5, i)
            day5.text=str(listm[5][i])
            
        if(count==4 or count==7):
            for i in range(7):
                table.rows[6].height=Inches(1)
                day6=table.cell(6,i)
                day6.text=str(listm[6][i])
            
        for n in range(6):
            cell=table.cell(n, 0)
            cell.text=""
            run = cell.paragraphs[0].add_run(str(listm[n][0]))
            run.font.color.rgb = RGBColor(255, 0, 0)
        count+=1
        doc.add_paragraph()
doc.add_picture(pic)
run=doc.add_heading("HOLIDAYS")
run.alignment = WD_ALIGN_PARAGRAPH.CENTER
table2=doc.add_table(rows=18, cols=3)
for i in range(18):
    h = holiday[i]
    run=table2.cell(i,0)
    run.text=h["date"]
for i in range(18):
    h=holiday[i]
    run=table2.cell(i,1)
    run.text=h["Day"]
for i in range(18):
    h=holiday[i]
    run=table2.cell(i,2)
    run.text=h["name"]
doc.save("calendar.docx")