from docx import Document
from docx.shared import Inches,RGBColor

list1=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",1,2,3],[4,5,6,7,8,9,10],[11,12,13,14,15,16,17],[18,19,20,21,22,23,24],[25,26,27,28,29,30,31]]
list2=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[1,2,3,4,5,6,7],[8,9,10,11,12,13,14],[15,16,17,18,19,20,21],[22,23,24,25,26,27,28],[".",".",".",".",".",".",".",]]
list3=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[1,2,3,4,5,6,7],[8,9,10,11,12,13,14],[15,16,17,18,19,20,21],[22,23,24,25,26,27,28],[29,30,31,".",".",".","."]]
list4=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",1,2,3,4],[5,6,7,8,9,10,11],[12,13,14,15,16,17,18],[19,20,21,22,23,24,25],[26,27,28,29,30,".","."]]
list5=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",".",1,2],[3,4,5,6,7,8,9],[10,11,12,13,14,15,16],[17,18,19,20,21,22,23],[24,25,26,27,28,29,30],[31,".",".",".",".",".","."]]
list6=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",1,2,3,4,5,6],[7,8,9,10,11,12,13],[14,15,16,17,18,19,20],[21,22,23,24,25,26,27],[28,29,30,".",".",".","."]]
list7=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",1,2,3,4],[5,6,7,8,9,10,11],[12,13,14,15,16,17,18],[19,20,21,22,23,24,25],[26,27,28,29,30,31,"."]]
list8=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",".",".",1],[2,3,4,5,6,7,8],[9,10,11,12,13,14,15],[16,17,18,19,20,21,22],[23,24,25,26,27,28,29],[30,31,".",".",".",".",","]]
list9=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",1,2,3,4,5],[6,7,8,9,10,11,12],[13,14,15,16,17,18,19],[20,21,22,23,24,25,26],[27,28,29,30,31,".","."]]
list10=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",1,2,3],[4,5,6,7,8,9,10],[11,12,13,14,15,16,17],[18,19,20,21,22,23,24],[25,26,27,28,29,30,31]]
list11=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[1,2,3,4,5,6,7],[8,9,10,11,12,13,14],[15,16,17,18,19,20,21],[22,23,24,25,26,27,28],[29,30,".",".",".",".","."]]
list12=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",1,2,3,4,5],[6,7,8,9,10,11,12],[13,14,15,16,17,18,19],[20,21,22,23,24,25,26],[27,28,29,30,31,".","."]]

listfull=[list1,list2,list3,list4,list5,list6,list7,list8,list9,list10,list11,list12]
year="2026"
month=["January","February","March","April","May","June","July","August","September","October","November","December"]
print("Done check")
count=0
doc=Document()
row=6
for listm in listfull:
    if(count<12):
        if(count==4 or count==7):
            row=7
        else:
            row=6
            
        table1=doc.add_table(rows=1,cols=3)

        mon=table1.cell(0,0)
        run=mon.paragraphs[0].add_run(str(month[count]))
        run.font.color.rgb=RGBColor(0,0,255)

        year1=table1.cell(0,1)
        run = year1.paragraphs[0].add_run(str(year))
        run.font.color.rgb = RGBColor(128, 0, 0)

        table = doc.add_table(rows=row,cols=7)    #<<
        table.style='Table Grid'

        table.columns[0].width, table.columns[1].width, table.columns[2].width, table.columns[3].width, table.columns[4].width, table.columns[5].width, table.columns[6].width =Inches(1.0), Inches(2.0), Inches(2.0), Inches(2.0), Inches(2.0), Inches(2.0), Inches(2.0), 
        table.rows[0].height, table.rows[1].height, table.rows[2].height, table.rows[3].height, table.rows[4].height, table.rows[5].height = Inches(0.5), Inches(1), Inches(1), Inches(1), Inches(1), Inches(1)

        for i in range(7):
            day0=table.cell(0, i)
            day0.text=str(listm[0][i])
            
        for i in range(7):
            day1=table.cell(1, i)
            day1.text=str(listm[1][i])
            
        for i in range(7):
            day2=table.cell(2, i)
            day2.text=str(listm[2][i])
            
        for i in range(7):
            day3=table.cell(3, i)
            day3.text=str(listm[3][i])

        for i in range(7):
            day4=table.cell(4, i)
            day4.text = str(listm[4][i])
            
        for i in range(7):
            day5=table.cell(5, i)
            day5.text=str(listm[5][i])
            
        if(count==4 or count==7):
            for i in range(7):
                table.rows[6].height=Inches(1)
                day6=table.cell(6,i)
                day6.text=str(listm[6][i])
            
        for n in range(6):
            cell=table.cell(n, 0)
            cell.text=""
            run = cell.paragraphs[0].add_run(str(listm[n][0]))
            run.font.color.rgb = RGBColor(255, 0, 0)
        count+=1
        
        
        
doc.save("calender.docx")