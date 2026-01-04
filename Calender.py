from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


list0=["Sun","Mon","Tue","Wed","Thu","Fri","Sat"]
list1=[".",".",".",1,2,3,4,]
list2=[5,6,7,8,9,10,11]
list3=[12,13,14,15,16,17,18]
list4=[19,20,21,22,23,24,25]
list5=[26,27,28,29,30,31,"."]
listm=[list0,list1,list2,list3,list4,list5]
year="2026"
month="jan"
listn=[year,listm]
print(listn)

doc=Document()

doc.add_heading("Calender")
doc.add_paragraph("2026 SMVITM")
table1=doc.add_table(rows=1,cols=3)

mon=table1.cell(0,0)
mon.text=month

year1=table1.cell(0,1)
year1.text=year

table = doc.add_table(rows=6,cols=7)
table.style='Table Grid'

table.columns[0].width =Inches(2.0) 
table.columns[1].width =Inches(2.0) 
table.columns[2].width =Inches(2.0) 
table.columns[3].width =Inches(2.0) 
table.columns[4].width =Inches(2.0) 
table.columns[5].width =Inches(2.0) 
table.columns[6].width =Inches(2.0)

table.rows[0].height =Inches(1) 
table.rows[1].height =Inches(1) 
table.rows[2].height =Inches(1) 
table.rows[3].height =Inches(1) 
table.rows[4].height =Inches(1) 
table.rows[5].height =Inches(1) 


for h in range(7):
    day0=table.cell(0, h)
    para=day0.add_paragraph('')
    run= para.add_run(str(list0[h]))

for i in range(7):
    day1=table.cell(1, i)
    para=day1.add_paragraph('')
    run= para.add_run(str(list1[i]))
    
for j in range(7):
    day2=table.cell(2, j)
    para=day2.add_paragraph('')
    run= para.add_run(str(list2[j]))

for k in range(7):
    day3=table.cell(3, k)
    para=day3.add_paragraph('')
    run= para.add_run(str(list3[k]))

for l in range(7):
    day4=table.cell(4, l)
    para=day4.add_paragraph('')
    run= para.add_run(str(list4[l]))
    
for m in range(7):
    day5=table.cell(5, m)
    para=day5.add_paragraph('')
    run= para.add_run(str(list5[m]))
    
for n in range(6):
    cell=table.cell(n, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'red') 
    tcPr.append(shd)
    
doc.save("calender.docx")